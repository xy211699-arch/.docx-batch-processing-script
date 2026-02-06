"""
Word文档批量处理工具
功能：字体统一格式化、插入动态页码、智能搜索并插入人物图片
"""

import os
import re
import requests
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import urllib.parse

# ==================== 配置模块 ====================

def get_user_desktop():
    """获取用户桌面路径（跨平台）"""
    return Path.home() / 'Desktop'

def create_output_folder(base_folder, folder_name):
    """创建输出文件夹"""
    folder_path = base_folder / folder_name
    folder_path.mkdir(exist_ok=True)
    return folder_path

# ==================== 文档格式化模块 ====================

def set_document_styles(doc):
    """
    统一设置文档样式：楷体四号加粗，调整页边距和行间距
    """
    # 设置正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '楷体'
    font.size = Pt(14)
    font.bold = True
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)      # 1英寸
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)     # 1.25英寸
        section.right_margin = Cm(3.18)
    
    # 处理所有段落文本
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '楷体'
            run.font.size = Pt(14)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        paragraph.paragraph_format.line_spacing = 1.5
    
    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '楷体'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def add_dynamic_page_numbers(doc):
    """
    在页脚居中插入动态页码（仅显示数字）
    """
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = footer_para.add_run()
        run.font.name = '楷体'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        # 创建动态页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

# ==================== 图片处理模块 ====================

def extract_name_from_document(doc):
    """
    从文档第一行提取人物姓名（支持中文姓名+英文名格式）
    格式示例：南部阳一郎（Yoichiro Nambu）或以撒·阿西莫夫（Isaac Asimov）
    """
    if len(doc.paragraphs) == 0:
        return None
    
    first_line = doc.paragraphs[0].text.strip()
    if not first_line:
        return None
    
    # 匹配括号前的中文姓名部分
    pattern = r'^([^（\(]+?)[（\(]'
    match = re.search(pattern, first_line)
    
    if match:
        chinese_name = match.group(1).strip()
        chinese_name = re.sub(r'\s+', ' ', chinese_name)
        return chinese_name
    
    # 如果没有括号格式，返回第一行内容
    return first_line[:50]  # 限制长度

def search_image_on_bing(person_name, search_headers=None):
    """
    使用必应搜索人物图片，返回图片URL
    """
    if not person_name:
        return None
    
    if search_headers is None:
        search_headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        }
    
    try:
        search_query = urllib.parse.quote(f"{person_name} 人物 照片")
        search_url = f"https://www.bing.com/images/search?q={search_query}&first=1"
        
        response = requests.get(search_url, headers=search_headers, timeout=15)
        if response.status_code != 200:
            return None
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 多种方式查找图片URL
        img_selectors = [
            ('img', {'class': 'mimg'}),
            ('img', {'src': re.compile(r'\.(jpg|jpeg|png|webp)', re.I)}),
            ('a', {'class': 'iusc'})
        ]
        
        for selector, attrs in img_selectors:
            elements = soup.find_all(selector, attrs)
            for element in elements:
                img_url = None
                
                if selector == 'img':
                    img_url = element.get('src') or element.get('data-src')
                elif selector == 'a':
                    m = re.search(r'murl":"([^"]+)"', str(element))
                    if m:
                        img_url = m.group(1).replace('\\/', '/')
                
                if img_url and img_url.startswith(('http://', 'https://')):
                    return img_url
                elif img_url and img_url.startswith('//'):
                    return 'https:' + img_url
        
        return None
        
    except Exception:
        return None

def download_image(image_url, save_path, download_headers=None):
    """
    下载图片到指定路径
    """
    if download_headers is None:
        download_headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
    
    try:
        response = requests.get(image_url, headers=download_headers, timeout=20, stream=True)
        if response.status_code == 200:
            with open(save_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            # 验证文件大小
            if os.path.getsize(save_path) > 2048:  # 至少2KB
                return True
            else:
                os.remove(save_path)
                return False
        return False
    except Exception:
        return False

def insert_image_into_document(doc, image_path):
    """
    在文档末尾插入图片，自动调整到合适大小
    """
    if not os.path.exists(image_path):
        return False
    
    try:
        # 添加分隔和标题
        doc.add_paragraph()
        title_para = doc.add_paragraph("人物图片：")
        title_run = title_para.runs[0]
        title_run.font.name = '宋体'
        title_run.font.size = Pt(12)
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建图片段落
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        
        # 计算合适的大小
        section = doc.sections[0]
        page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
        
        try:
            from PIL import Image
            with Image.open(image_path) as img:
                img_width_px, img_height_px = img.size
                dpi = img.info.get('dpi', (96, 96))[0]
                
                img_width_cm = img_width_px * 2.54 / dpi
                max_width_cm = page_width_cm * 0.7
                
                if img_width_cm > max_width_cm:
                    scale_ratio = max_width_cm / img_width_cm
                    target_width_cm = max_width_cm
                elif img_width_cm < 5:
                    target_width_cm = 5
                else:
                    target_width_cm = img_width_cm
                
                run.add_picture(image_path, width=Cm(target_width_cm))
        except (ImportError, Exception):
            # Pillow不可用或出错时使用默认大小
            default_width_cm = page_width_cm * 0.5
            run.add_picture(image_path, width=Cm(default_width_cm))
        
        doc.add_paragraph()
        return True
        
    except Exception:
        return False

# ==================== 文件处理模块 ====================

def backup_documents(source_folder, backup_name=None):
    """
    创建文档备份
    """
    if backup_name is None:
        backup_name = f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    backup_folder = source_folder / backup_name
    backup_folder.mkdir(exist_ok=True)
    
    import shutil
    for file in source_folder.glob("*.docx"):
        shutil.copy2(file, backup_folder / file.name)
    
    return backup_folder

def process_document_collection(folder_path, image_storage_path, enable_backup=True):
    """
    批量处理文档集合
    返回处理统计和错误记录
    """
    if not folder_path.exists():
        return {'error': '文件夹不存在', 'failed': []}
    
    # 创建备份（可选）
    backup_path = None
    if enable_backup:
        backup_path = backup_documents(folder_path)
    
    # 初始化统计
    statistics = {
        'total': 0,
        'success': 0,
        'images_found': 0,
        'images_inserted': 0,
        'backup': backup_path,
        'failed': []
    }
    
    # 处理每个文档
    for doc_file in folder_path.glob("*.docx"):
        statistics['total'] += 1
        file_record = {
            'filename': doc_file.name,
            'errors': [],
            'chinese_name': None
        }
        
        try:
            # 1. 打开文档
            doc = Document(doc_file)
            
            # 2. 格式化文档
            set_document_styles(doc)
            add_dynamic_page_numbers(doc)
            
            # 3. 提取姓名并处理图片
            person_name = extract_name_from_document(doc)
            if person_name:
                file_record['chinese_name'] = person_name
                
                # 搜索并下载图片
                image_url = search_image_on_bing(person_name)
                if image_url:
                    safe_name = re.sub(r'[^\w\u4e00-\u9fa5\-]', '_', person_name)
                    image_filename = f"{safe_name}.jpg"
                    image_path = image_storage_path / image_filename
                    
                    if download_image(image_url, image_path):
                        statistics['images_found'] += 1
                        
                        # 插入图片到文档
                        if insert_image_into_document(doc, image_path):
                            statistics['images_inserted'] += 1
            
            # 4. 保存文档
            doc.save(doc_file)
            statistics['success'] += 1
            
        except Exception as e:
            file_record['errors'].append(f"处理异常: {str(e)}")
            statistics['failed'].append(file_record)
    
    return statistics

def generate_error_report(failed_records, report_path):
    """
    生成错误报告文件
    """
    if not failed_records:
        return None
    
    try:
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(f"Word文档批量处理错误报告\n")
            f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"失败文件数: {len(failed_records)}\n")
            f.write("=" * 60 + "\n\n")
            
            for i, record in enumerate(failed_records, 1):
                f.write(f"{i}. 文件名: {record['filename']}\n")
                if record.get('chinese_name'):
                    f.write(f"   人物姓名: {record['chinese_name']}\n")
                if record.get('errors'):
                    f.write(f"   错误信息:\n")
                    for error in record['errors']:
                        f.write(f"     - {error}\n")
                f.write("\n")
        
        return report_path
    except Exception:
        return None

# ==================== 主程序模块 ====================



def main():
    """
    主函数：用户交互和流程控制
    """
    print("=" * 50)
    print("Word文档批量处理工具")
    print("功能：格式化文档 | 插入页码 | 智能添加人物图片")
    print("=" * 50)
    
    # 获取用户输入
    desktop = get_user_desktop()
    
    try:
        folder_name = input("请输入要处理的文件夹名称: ").strip()
        source_folder = desktop / folder_name
        
        if not source_folder.exists():
            print(f"错误：文件夹 '{folder_name}' 不存在！")
            return
        
        # 创建图片存储文件夹
        images_folder = create_output_folder(desktop, "人物图片")
        
        # 确认操作
        print(f"\n处理配置：")
        print(f"  源文件夹: {source_folder}")
        print(f"  图片保存: {images_folder}")
        print(f"  文档格式: 楷体四号加粗 + 居中页码")
        
        confirm = input("\n确认开始处理？(y/n): ").strip().lower()
        if confirm != 'y':
            print("操作已取消。")
            return
        
        # 执行处理
        print("\n开始处理文档，请稍候...")
        print("=" * 50)
        
        results = process_document_collection(source_folder, images_folder, enable_backup=True)
        
        # 显示结果
        print("\n" + "=" * 50)
        print("处理完成！")
        print(f"总计文档: {results['total']}")
        print(f"成功处理: {results['success']}")
        print(f"找到图片: {results['images_found']}")
        print(f"插入图片: {results['images_inserted']}")
        
        if results.get('backup'):
            print(f"文档备份: {results['backup']}")
        
        # 生成错误报告（如果有）
        if results['failed']:
            report_file = desktop / "文档处理错误报告.txt"
            report_path = generate_error_report(results['failed'], report_file)
            if report_path:
                print(f"错误报告: {report_path}")
        
        print("=" * 50)
        
    except KeyboardInterrupt:
        print("\n\n用户中断操作。")
    except Exception as e:
        print(f"\n程序执行出错: {e}")

# ==================== 程序入口 ====================

if __name__ == "__main__":
    main()